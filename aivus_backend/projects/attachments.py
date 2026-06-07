"""Shared limits and helpers for brief attachments.

Kept dependency-free of the views/tasks modules so both can import it without
creating an import cycle (views imports tasks, tasks imports this).
"""

from __future__ import annotations

import io
import logging
import urllib.request
from urllib.parse import urlparse

import magic
from docx import Document

logger = logging.getLogger(__name__)

MAX_ATTACHMENT_SIZE_BYTES = 10 * 1024 * 1024
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "image/jpeg",
    "image/png",
    "image/webp",
    "image/gif",
    "text/plain",
    DOCX_MIME,
}

# libmagic reports a .docx (a zip container) inconsistently across versions, so
# accept the zip/octet-stream sniff results when the client declared a docx.
DOCX_DETECTED_ALIASES = {
    DOCX_MIME,
    "application/zip",
    "application/x-zip-compressed",
    "application/octet-stream",
}

DOCX_MAX_TEXT_CHARS = 100_000

REMOTE_DOWNLOAD_TIMEOUT_SECONDS = 15

WIX_FILE_HOST_SUFFIXES = (
    "wixstatic.com",
    "wixmp.com",
    "wix.com",
    "wixsite.com",
    "usrfiles.com",
)


class _NoRedirectHandler(urllib.request.HTTPRedirectHandler):
    # SSRF guard: never follow redirects, otherwise a whitelisted host could
    # bounce the request to an internal address.
    def redirect_request(self, *args, **kwargs):
        return None


def sniff_mime(fileobj) -> str:
    """Detect MIME type from file contents via libmagic, not the client-declared
    Content-Type. Rewinds the file pointer so the caller can still read it."""
    try:
        sample = fileobj.read(4096)
    finally:
        try:
            fileobj.seek(0)
        except Exception:
            logger.exception("Cannot rewind uploaded file")
    try:
        return magic.from_buffer(sample, mime=True) or ""
    except Exception:
        logger.exception("magic.from_buffer failed")
        return ""


def sniff_mime_bytes(data: bytes) -> str:
    try:
        return magic.from_buffer(data[:4096], mime=True) or ""
    except Exception:
        logger.exception("magic.from_buffer failed")
        return ""


def extract_docx_text(data: bytes) -> str:
    """Flatten a .docx file to plain text (paragraphs and table cells).

    Gemini cannot read Word documents directly, so the content is extracted and
    passed to the model as text. Returns an empty string on any failure.
    """
    try:
        document = Document(io.BytesIO(data))
    except Exception:
        logger.exception("Cannot parse docx attachment")
        return ""

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            joined = " | ".join(cell for cell in cells if cell)
            if joined:
                lines.append(joined)

    return "\n".join(lines)[:DOCX_MAX_TEXT_CHARS]


def download_remote_file(
    url: str,
    *,
    allowed_host_suffixes: tuple[str, ...],
    max_bytes: int = MAX_ATTACHMENT_SIZE_BYTES,
) -> tuple[bytes, str] | None:
    """Download a remote file with SSRF guards, a size cap and MIME sniffing.

    Returns (data, mime) when the file is fetched and its sniffed MIME is in the
    allow-list, otherwise None. Only http(s) URLs whose host matches one of
    allowed_host_suffixes are fetched; redirects are not followed.
    """
    if not isinstance(url, str) or not url:
        return None

    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        logger.warning("Rejected non-http(s) attachment url")
        return None

    host = (parsed.hostname or "").lower()
    if not host or not any(
        host == suffix or host.endswith("." + suffix)
        for suffix in allowed_host_suffixes
    ):
        logger.warning("Rejected attachment url with disallowed host: %s", host)
        return None

    opener = urllib.request.build_opener(_NoRedirectHandler)
    request = urllib.request.Request(  # noqa: S310 scheme validated above
        url, headers={"User-Agent": "aivus-brief-intake"}
    )
    try:
        with opener.open(request, timeout=REMOTE_DOWNLOAD_TIMEOUT_SECONDS) as response:
            declared_length = response.headers.get("Content-Length")
            if (
                declared_length
                and declared_length.isdigit()
                and int(declared_length) > max_bytes
            ):
                logger.warning("Attachment too large by header: %s", host)
                return None
            data = response.read(max_bytes + 1)
    except Exception:
        logger.exception("Attachment download failed")
        return None

    if not data or len(data) > max_bytes:
        logger.warning("Attachment empty or exceeded max size: %s", host)
        return None

    mime = sniff_mime_bytes(data).lower()
    if mime not in ALLOWED_MIME_TYPES:
        logger.warning("Attachment MIME not allowed: %s", mime)
        return None

    return data, mime
